#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修复Word文件中的星号显示问题
将错误的星号格式转换为正确的格式
"""

from docx import Document
import os


def convert_stars_to_correct_format(text):
    """
    将错误的星号格式转换为正确的格式
    
    错误格式：
    - "3星★(4星)★(5星)" 应该是 "*****(5星)"
    - "3星★(4星)" 应该是 "****(4星)"
    - "3星" 应该是 "***(3星)"
    
    正确格式：用*表示星级，后面跟括号标注
    """
    # 处理 "3星★(4星)★(5星)" -> "*****(5星)"
    if "3星★(4星)★(5星)" in text:
        text = text.replace("3星★(4星)★(5星)", "*****(5星)")
    
    # 处理 "3星★(4星)" -> "****(4星)"
    if "3星★(4星)" in text:
        text = text.replace("3星★(4星)", "****(4星)")
    
    # 处理 "3星" -> "***(3星)"
    # 注意：需要确保不是"3星★"的情况
    if text.strip().endswith("3星") or " 3星 " in text or text.strip() == "3星":
        text = text.replace("3星", "***(3星)")
    
    return text


def fix_word_file(input_file, output_file=None):
    """
    修复Word文件中的星号显示
    
    Args:
        input_file: 输入Word文件路径
        output_file: 输出Word文件路径（如果为None，则覆盖原文件）
    """
    if output_file is None:
        output_file = input_file
    
    print(f"正在修复文件: {input_file}")
    
    # 读取Word文件
    doc = Document(input_file)
    
    fix_count = 0
    
    # 修复表格中的内容
    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                fixed_text = convert_stars_to_correct_format(original_text)
                
                if fixed_text != original_text:
                    # 清空单元格并重新写入
                    cell.text = ""
                    for paragraph in cell.paragraphs:
                        # 保留原有的格式
                        if paragraph.runs:
                            # 如果有run，保持原有格式
                            run = paragraph.runs[0]
                            run.text = fixed_text
                        else:
                            # 如果没有run，直接写入
                            paragraph.text = fixed_text
                    
                    fix_count += 1
                    print(f"  修复: '{original_text}' -> '{fixed_text}'")
    
    # 修复正文段落
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        fixed_text = convert_stars_to_correct_format(original_text)
        
        if fixed_text != original_text:
            # 清空段落并重新写入
            for run in paragraph.runs:
                run.text = run.text.replace("3星★(4星)★(5星)", "*****(5星)")
                run.text = run.text.replace("3星★(4星)", "****(4星)")
                if run.text.strip().endswith("3星"):
                    run.text = run.text.replace("3星", "***(3星)")
            
            fix_count += 1
            print(f"  修复: '{original_text}' -> '{fixed_text}'")
    
    # 保存文件
    doc.save(output_file)
    
    print(f"\n修复完成！共修复 {fix_count} 处")
    print(f"输出文件: {output_file}")
    
    return fix_count


def main():
    """主函数"""
    # 要修复的文件列表
    files_to_fix = [
        "output/武侯区供电中心服务提升应用建议_详尽版.docx",
        "output/武侯区供电中心服务提升应用建议.docx",
    ]
    
    total_fixes = 0
    
    for file_path in files_to_fix:
        if os.path.exists(file_path):
            fix_count = fix_word_file(file_path)
            total_fixes += fix_count
            print()
        else:
            print(f"文件不存在: {file_path}\n")
    
    print(f"\n总共修复了 {total_fixes} 处星号显示问题")


if __name__ == "__main__":
    main()
